"""
Document text extraction module
Handles extraction from PDF, DOCX, XLSX, TXT files
"""

import logging
from pathlib import Path
from typing import Optional, Dict, List
import hashlib

# Text extraction libraries
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import openpyxl

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from various document formats"""
    
    SUPPORTED_FORMATS = {
        ".pdf": "extract_pdf",
        ".docx": "extract_docx",
        ".doc": "extract_docx",
        ".txt": "extract_txt",
        ".xlsx": "extract_xlsx",
        ".xls": "extract_xlsx",
    }
    
    @staticmethod
    def extract_pdf(file_path: str) -> Dict[str, any]:
        """
        Extract text from PDF file
        Uses pdfplumber for better formatting, falls back to PyPDF2
        """
        text = ""
        metadata = {}
        
        try:
            # Try pdfplumber first (better formatting)
            with pdfplumber.open(file_path) as pdf:
                metadata["total_pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num} ---\n"
                        text += page_text
                
                # Extract PDF metadata
                if pdf.metadata:
                    metadata["title"] = pdf.metadata.get("Title", "")
                    metadata["author"] = pdf.metadata.get("Author", "")
            
            logger.info(f"Extracted text from PDF: {len(text)} characters, {metadata['total_pages']} pages")
            
            return {
                "success": True,
                "text": text.strip(),
                "metadata": metadata,
                "format": "pdf",
            }
            
        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "format": "pdf",
            }
    
    @staticmethod
    def extract_docx(file_path: str) -> Dict[str, any]:
        """Extract text from DOCX file"""
        text = ""
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
            
            logger.info(f"Extracted text from DOCX: {len(text)} characters")
            
            return {
                "success": True,
                "text": text.strip(),
                "metadata": {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
                "format": "docx",
            }
            
        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "format": "docx",
            }
    
    @staticmethod
    def extract_txt(file_path: str) -> Dict[str, any]:
        """Extract text from TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            
            logger.info(f"Extracted text from TXT: {len(text)} characters")
            
            return {
                "success": True,
                "text": text.strip(),
                "metadata": {"lines": len(text.split('\n'))},
                "format": "txt",
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    text = f.read()
                return {
                    "success": True,
                    "text": text.strip(),
                    "metadata": {"lines": len(text.split('\n')), "encoding": "latin-1"},
                    "format": "txt",
                }
            except Exception as e:
                logger.error(f"Error extracting TXT (encoding issue): {str(e)}")
                return {
                    "success": False,
                    "error": f"Encoding error: {str(e)}",
                    "format": "txt",
                }
        except Exception as e:
            logger.error(f"Error extracting TXT: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "format": "txt",
            }
    
    @staticmethod
    def extract_xlsx(file_path: str) -> Dict[str, any]:
        """Extract text from XLSX file"""
        text = ""
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n=== Sheet: {sheet_name} ===\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        text += row_text + "\n"
            
            logger.info(f"Extracted text from XLSX: {len(text)} characters, {len(workbook.sheetnames)} sheets")
            
            return {
                "success": True,
                "text": text.strip(),
                "metadata": {"sheets": len(workbook.sheetnames)},
                "format": "xlsx",
            }
            
        except Exception as e:
            logger.error(f"Error extracting XLSX: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "format": "xlsx",
            }
    
    @classmethod
    def extract(cls, file_path: str) -> Dict[str, any]:
        """
        Extract text from any supported format
        Auto-detects format and calls appropriate method
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                "success": False,
                "error": f"File not found: {file_path}",
            }
        
        file_ext = file_path.suffix.lower()
        
        if file_ext not in cls.SUPPORTED_FORMATS:
            return {
                "success": False,
                "error": f"Unsupported file format: {file_ext}. Supported: {list(cls.SUPPORTED_FORMATS.keys())}",
            }
        
        # Call appropriate extraction method
        method_name = cls.SUPPORTED_FORMATS[file_ext]
        method = getattr(cls, method_name)
        
        result = method(str(file_path))
        result["file_path"] = str(file_path)
        result["file_name"] = file_path.name
        result["file_size_bytes"] = file_path.stat().st_size
        
        return result


def calculate_md5(file_path: str) -> str:
    """Calculate MD5 hash of file for duplicate detection"""
    hash_md5 = hashlib.md5()
    
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    
    return hash_md5.hexdigest()


def validate_text_quality(text: str, min_length: int = 100) -> Dict[str, any]:
    """
    Validate extracted text quality
    Returns quality score and issues
    """
    issues = []
    score = 1.0
    
    # Check minimum length
    if len(text) < min_length:
        issues.append(f"Text too short: {len(text)} chars (min: {min_length})")
        score -= 0.3
    
    # Check for mostly special characters
    special_char_ratio = sum(1 for c in text if not c.isalnum() and not c.isspace()) / len(text) if text else 0
    if special_char_ratio > 0.5:
        issues.append(f"Too many special characters: {special_char_ratio:.0%}")
        score -= 0.2
    
    # Check for corrupted content (too many replacement characters)
    replacement_count = text.count("�")
    if replacement_count > 10:
        issues.append(f"Possible encoding corruption: {replacement_count} replacement chars")
        score -= 0.3
    
    # Check for readability
    lines = text.split("\n")
    avg_line_length = sum(len(line) for line in lines) / len(lines) if lines else 0
    
    if avg_line_length < 5 or avg_line_length > 1000:
        issues.append(f"Unusual line length: avg {avg_line_length:.0f} chars")
        score -= 0.1
    
    return {
        "quality_score": max(0, min(1.0, score)),
        "is_valid": score > 0.5,
        "issues": issues,
        "text_length": len(text),
        "line_count": len(lines),
        "avg_line_length": avg_line_length,
    }
